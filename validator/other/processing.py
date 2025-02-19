import docx

class Linking():
    pass

class Word():
    def __init__(self):
        pass

    def findReplace(self, file, lets, savePath=False):
        """ Функция производит поиск в абзацах и таблицах производя замену.
            
            Keyword arguments:
                file -> str path
                lets -> [{"var":"", "val":""}]
            
        """
        doc = docx.Document(file)
        for paragraph in (doc.paragraphs):
            for let in lets:
                if paragraph.text.count(let['var']):
                    paragraph.text = paragraph.text.replace(let['var'], let['val'])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for let in lets:
                        if cell.text.count(let['var']):
                            cell.text = cell.text.replace(let['var'], let['val'])
        if savePath:
            doc.save(savePath)
        else:
            doc.save(file)

    def setStyle(self):
        pass

class Excel():
    pass